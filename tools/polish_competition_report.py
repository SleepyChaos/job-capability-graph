from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"D:\git\job-capability-graph\outputs\institution_database_report\具身智能机构库与企业产业链数据建设报告_比赛版_20260825.docx")
OUTPUT = Path(r"D:\git\job-capability-graph\outputs\institution_database_report\具身智能机构库与企业产业链数据建设报告_比赛润色版_20260825.docx")


PARAGRAPH_REPLACEMENTS = {
    2: "字段体系 · 产业链分类 · 数据核验 · 更新方案",
    3: "数据依据：机构主库与企业补充数据表（2026年8月25日批次）",
    4: "统计范围：机构、企业产业链、人才、成果与技术方向数据",
    9: "具身智能产业涉及材料、传感器、执行系统、人工智能模型、机器人本体、场景方案以及政策生态等多类主体。仅维护企业名称清单，既无法说明机构在产业链中的位置，也难以关联人才、成果和技术方向。因此，项目以机构为统一主体，将企业、高校、科研院所和政府/事业单位纳入同一机构库，并通过统一名称、产业链标签、产品信息、经营状态以及成果和人才关系，形成可检索、可统计、可持续更新的数据基础。",
    10: "本次整理以机构主库统一管理机构数据，并使用企业补充数据表补充产业链和企业经营信息。数据处理遵循三项原则：一是机构主库用于汇总机构信息，企业补充数据表是产业链分类的数据来源之一；二是分类结论须由机构业务、产品和公开证据共同支持；三是对重复、缺失或冲突数据不作强行覆盖，而是保留原始记录并转入人工复核。",
    18: "经名称核验，企业数据表的656行名称均能在机构库代表名称中找到。需要说明的是，656行并不等同于656个独立实体：表内存在4组完全重复名称；名称规范化后，还存在1组仅中英文括号不同的记录。因此，本报告统一采用“656行企业分类数据、652个不同原始名称”的表述。",
    23: "（一）35个字段的分类说明",
    24: "机构库共设置35个字段。为便于评审理解数据结构和字段用途，本报告按照字段在数据建设中的作用，将其归纳为主体基础、地理与来源、产业与产品、运营发展、成果与能力五类。该分类仅用于说明字段之间的业务关系，不改变原有字段名称和含义。",
    26: "图2  机构库35个字段的分类及其关系",
    27: "五类字段分别用于回答以下问题：主体基础字段用于确认机构身份；地理与来源字段用于记录机构所在地和信息依据；产业与产品字段用于说明机构所处产业链环节及其产品或服务；运营发展字段用于记录机构的发展状态；成果与能力字段用于说明能够反映其技术积累的外部证据。下表仅汇总字段数量和代表字段，具体含义在表后分别说明。",
    30: "1. 主体基础字段：机构识别与名称统一",
    32: "这类字段不直接评价机构能力，而是保证后续产品、人才、专利和技术信息能够关联到正确主体。若机构名称合并错误，后续产业链统计和机构分析都会受到影响，因此主体识别与名称统一是其他数据处理工作的前提。",
    33: "2. 地理与来源字段：空间定位与信息追溯",
    35: "这类字段将机构所在地和信息来源分别记录。地理字段用于区域统计，来源和校验字段用于事实核验，避免将无法追溯的信息直接写入正式结论。",
    36: "3. 产业与产品字段：产业位置与产品能力说明",
    38: "上述六个字段共同构成分类依据。产业链类别和层级给出分类结果，细分领域和代表产品说明分类所对应的业务，产品类型和关键特性说明作出判断的依据，从而避免仅根据企业名称或单个关键词分类。",
    39: "4. 运营发展字段：产业化状态与经营信息",
    41: "这类字段用于判断机构处于技术验证、产品交付还是规模化扩展阶段。融资信息只是反映发展状态的辅助证据，不能代替产品和量产证据，也不能单独决定产业链层级。",
    42: "5. 成果与能力字段：技术积累的外部证据",
    44: "这类字段不只依据企业自述，而是通过专利、标准、人才和技术标注提供可回查的外部证据。L2用于概括机构的主要技术方向，L3用于识别更具体的技术内容；两级标注均保留明细记录，以便核验和解释。",
    47: "（一）分类框架的形成过程",
    48: "产业链层级并非来源表中直接给出的结论，而是项目组根据机构业务和产品证据进行整理。分类时，先识别机构主营业务和代表产品，再结合产品类型、关键技术特征与细分领域确定其主要产业位置；随后参考量产、运营、融资和招聘信息判断其发展状态，并通过官网、官方招聘、公告等公开来源交叉核验。",
    56: "具身智能本体承担部件、控制系统和任务能力的综合集成，是连接基础供给与场景应用的核心载体，因此归入中游；软件与算法、AI大模型和决策认知系统负责感知理解、规划决策、控制和学习，为本体提供主要的智能功能，因此主要归入中游。AI大模型中有2条记录被标为横向支撑，原因是相关机构更侧重为不同机构和场景提供通用平台服务，而非提供某一类本体产品的内部模块。",
    57: "场景应用与解决方案直接面向工业、医疗、家庭、商业服务等需求，完成系统部署和业务交付，因此归入下游；政策与生态主要提供政策组织、产业协同、公共平台和相关服务，不属于单一产品供给环节，因此归入横向支撑。由于工作簿未提供第12类的代码表，本报告不推测其名称和层级。",
    71: "下表集中展示四个层级的记录数量、占比和主要内容，用于核对前述分析。表中的层级说明是本项目采用的分类口径，不代表对机构技术水平高低的排序。",
    83: "机构库不仅保存机构和产品信息，还通过关系表关联人才、成果和技术方向。当前人才库有7,815条记录，人才机构成果关系7,603条，机构标准关系7,237条，成果技术标注7,586条。这些关系为分析机构的人才、成果和技术情况提供了企业自述以外的证据。",
    84: "专利族用于反映机构形成技术成果的情况，标准起草记录用于反映机构参与行业规范建设的情况，关联人才用于说明相关技术活动所涉及的人员和团队。成果技术标注进一步将专利、标准等成果对应到L2和L3技术方向，使分析从成果数量延伸到成果所属的技术方向。",
    85: "L2技术方向用于概括机构的主要技术领域，L3技术方向用于识别更具体的研究或产品方向。两级技术标注均应结合成果原文和映射记录使用，不能只依据标签数量判断机构技术能力。下表汇总当前数据规模和用途，为前述关系说明提供数量依据。",
    87: "当前L2技术方向覆盖率为72.81%，L3技术方向覆盖率为68.91%。对于尚未形成技术标注的机构，应进一步区分无可用成果、成果尚未提取、成果尚未映射以及映射需要人工复核等情况，不能将空白记录直接解释为机构不具备相关技术能力。",
    88: "八、动态更新方案与信息核验",
    89: "（一）公司、人员、成果、技术与事件的关联方式",
    90: "下一阶段拟按照“公司—人员—成果对象—技术—事件—新增主体”的关联顺序更新数据。首先，从机构官网、公告、招聘、专利、标准和政府项目中识别公司变化；再核对创始人、高管、技术负责人、发明人、标准起草人和论文作者；随后使用专利申请号、标准编号、DOI、产品型号或项目编号关联成果；最后将成果对应到现有L2/L3技术方向，并从产品发布、融资、获奖、展会、合作、招投标和招聘等事件中发现新的公司、人员和成果线索。",
    92: "（二）动态更新流程",
    94: "暂存：记录来源URL、页面标题、发布者、发布时间、采集时间以及支持相关事实的原文。",
    95: "名称核对与关联：分别核对公司、人员、产品、技术和事件名称；名称相同但证据不足时，不直接合并记录。",
    96: "数据比对：将候选事实与库内已有记录逐项核对，识别新增、更新、失效和冲突，并保留变更前内容。",
    97: "人工复核：主体、产业链和人员关系等需要判断的字段由人工确认；格式规则明确的字段可按既定规则处理。",
    98: "数据发布：形成经审核、可追溯的记录，注明复核人员、复核时间和修改原因，不直接覆盖原始记录。",
    100: "为避免将搜索结果直接作为事实依据，项目将数据来源分为三级。一级来源用于确认主体、专利、标准、公告和政府项目；二级来源用于补充产品、人员、实验室和活动信息；三级来源仅用于发现线索，相关信息须回到官方或可核验页面复核。",
    102: "九、数据准确性核验方案",
    103: "数据核验以可回查的原始材料为依据。机构主体核对名称、曾用名和官网；产业链分类核对代表产品、主营业务与层级规则；人员关系核对机构、职务和任职时间；成果对象通过申请号、标准号、DOI、型号或项目号回查；技术对应关系保留原文、L2/L3编码和规则版本。",
    105: "后续可从各类关系中抽取人工标注样本，计算准确率（Precision）、召回率（Recall）、F1值和混淆矩阵；同时统计来源证据覆盖率、时效合格率、冲突检出率和人工复核通过率。当前工作簿尚未提供完整的人工标注样本，因此本报告不列示未经验证的准确率指标。",
    106: "十、存在的问题及下一步工作",
    112: "建立独立事件暂存表和版本发布记录，使动态搜索结果能够进入发现、复核、发布和追溯流程。",
    113: "建立人工标注样本后再报告准确率，不用主观抽查结果替代准确率（Precision）、召回率（Recall）和F1值。",
    116: "本报告以2026年8月25日同批次工作簿为统计依据，文中数量由该批次工作簿重新统计形成。",
    117: "“产业链（12类标准）”沿用工作簿字段名称；当前批次仅有11类实际数据，不补写未出现的类别。",
    118: "动态更新、独立事件表和准确率实验属于下一阶段工作，不作为当前已完成成果。",
}


def replace_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_callout(table, label: str, body: str) -> None:
    paragraph = table.cell(0, 0).paragraphs[0]
    if len(paragraph.runs) >= 2:
        paragraph.runs[0].text = label + "  "
        paragraph.runs[1].text = body
        for run in paragraph.runs[2:]:
            run.text = ""
    else:
        replace_paragraph(paragraph, label + "  " + body)


def replace_cell_text(cell, text: str) -> None:
    replace_paragraph(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph(paragraph, "")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def replace_field_diagram(doc: Document) -> None:
    anchor = doc.paragraphs[25]
    replace_paragraph(anchor, "机构库35个字段分类")
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor.paragraph_format.space_before = Pt(2)
    anchor.paragraph_format.space_after = Pt(4)
    anchor.runs[0].bold = True
    anchor.runs[0].font.size = Pt(11)
    anchor.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    data = [
        ("主体基础", "7个字段", "机构ID、归一键、代表名称、机构类型、英文名/别名、原始形态、类型调整留痕", "183B66"),
        ("地理与来源", "9个字段", "数据来源、官网、总部城市、国家/省/市/区县、企业库关联、校验备注", "2F75B5"),
        ("产业与产品", "6个字段", "产业链类别、层级、细分领域、代表产品、产品类型、关键特性/参数", "188A82"),
        ("运营发展", "5个字段", "招聘链接、量产进展、运营路径、融资阶段、融资轮次分类", "A36A00"),
        ("成果与能力", "8个字段", "专利、标准、关联人才、任职高校人才、L2/L3技术方向及明细", "7B55B4"),
    ]
    table = doc.add_table(rows=len(data), cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.45), Inches(0.85), Inches(4.20)]
    for row, (label, count, fields, color) in zip(table.rows, data):
        row.cells[0].width, row.cells[1].width, row.cells[2].width = widths
        values = (label, count, fields)
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            replace_cell_text(cell, value)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.runs[0]
            run.font.size = Pt(8.5 if index == 2 else 9)
            if index < 2:
                run.bold = True
            if index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, color)
            elif index == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = RGBColor(31, 78, 121)
                set_cell_shading(cell, "EAF0F6")
            else:
                set_cell_shading(cell, "F7F9FB")
    anchor._p.addnext(table._tbl)


def main() -> None:
    doc = Document(SOURCE)
    for index, text in PARAGRAPH_REPLACEMENTS.items():
        replace_paragraph(doc.paragraphs[index], text)

    replace_callout(
        doc.tables[0],
        "建设目标",
        "围绕具身智能机构、企业产业链、人才、成果和技术方向数据，说明数据来源、字段设置、分类方法、核验要求及后续更新方案。",
    )
    replace_callout(
        doc.tables[1],
        "现有成果",
        "截至2026年8月25日，机构库已收录1,618条机构记录，设置35个字段；其中652条记录已完成产业链类别和层级标注。",
    )

    # 将自定义的“分层”名称改为普通的字段分类表述。
    table = doc.tables[3]
    replace_cell_text(table.cell(0, 0), "字段类别")
    for row, label in enumerate(["主体基础", "地理与来源", "产业与产品", "运营发展", "成果与能力"], start=1):
        replace_cell_text(table.cell(row, 0), label)

    replace_callout(
        doc.tables[4],
        "口径说明",
        "产业链等企业补充字段的完整度以652条已分类记录为分母。其余机构的相关字段为空，仅表示尚未补充，不能据此认定该机构不具备相应产品、技术或经营活动。",
    )
    replace_callout(
        doc.tables[5],
        "分类说明",
        "工作簿字段名称为“12类标准”，但当前批次仅实际出现11个类别，且未随表提供完整的12类代码表。因此，本报告仅统计当前出现的11类，不推测或补写缺失类别。",
    )

    # 删除“画像”等包装性说法，改用直接、可核验的用途说明。
    replace_cell_text(doc.tables[10].cell(4, 2), "汇总机构的主要技术方向")
    replace_cell_text(doc.tables[10].cell(5, 2), "识别更具体的技术方向")
    replace_callout(
        doc.tables[11],
        "信息关联方式",
        "公司 → 人员 → 产品/专利/标准/论文/项目 → L2/L3技术方向 → 发布、融资、获奖、合作、招聘等事件 → 新公司、新人员或新成果 → 主体核验。",
    )
    replace_callout(
        doc.tables[14],
        "项目建设成效",
        "项目已形成包含1,618条机构记录和35个字段的具身智能机构库，将企业、高校、科研院所和政府/事业单位统一纳入机构库管理。针对企业数据，项目依据主营业务、代表产品、技术特征和产业化信息，按照上游、中游、下游和横向支撑进行产业链分类，并通过官网、招聘、公告等来源交叉核验。目前652条记录已形成产业链标签，其中上游124条、中游437条、下游52条、横向支撑39条。对重复、缺失和冲突数据，保留原值并转入人工复核，不作强行补全。",
    )
    replace_field_diagram(doc)

    # 避免文件属性中继续出现“重构稿”等过程性信息。
    doc.core_properties.title = "具身智能机构库与企业产业链数据建设报告"
    doc.core_properties.subject = "比赛文稿"
    doc.core_properties.comments = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
