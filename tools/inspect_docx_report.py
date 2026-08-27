import argparse
import json
from pathlib import Path
from docx import Document


SOURCE = Path(r"D:\git\job-capability-graph\outputs\institution_database_report\具身智能机构库与企业产业链数据建设报告_比赛版_20260825.docx")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=SOURCE)
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int)
    parser.add_argument("--tables", action="store_true")
    parser.add_argument("--table-index", type=int, action="append")
    parser.add_argument("--runs", action="store_true")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    doc = Document(args.input)
    end = args.end if args.end is not None else len(doc.paragraphs)
    if args.json:
        payload = {
            "paragraphs": [
                {"i": i, "style": p.style.name, "text": p.text,
                 "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic, "size": r.font.size.pt if r.font.size else None} for r in p.runs] if args.runs else []}
                for i, p in enumerate(doc.paragraphs[args.start:end], start=args.start)
                if p.text.strip()
            ],
            "tables": [
                [[{"text": cell.text, "paragraphs": [
                    {"text": p.text, "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic, "size": r.font.size.pt if r.font.size else None} for r in p.runs]}
                    for p in cell.paragraphs
                ] if args.runs else []} for cell in row.cells] for row in table.rows]
                for ti, table in enumerate(doc.tables)
                if args.tables and (not args.table_index or ti in args.table_index)
            ],
        }
        print(json.dumps(payload, ensure_ascii=True))
        return
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)} sections={len(doc.sections)}")
    print("\n=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs[args.start:end], start=args.start):
        text = p.text.replace("\t", "\\t").replace("\n", "\\n")
        if text.strip():
            print(f"P{i:04d}\t[{p.style.name}]\t{text}")
    print("\n=== TABLES ===")
    for ti, table in enumerate(doc.tables if args.tables else []):
        print(f"TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else ''}")
        for ri, row in enumerate(table.rows):
            cells = [c.text.replace("\t", " ").replace("\n", " / ") for c in row.cells]
            print(f"T{ti}R{ri}\t" + "\t".join(cells))
    print("\n=== SECTIONS ===")
    for i, s in enumerate(doc.sections):
        print(i, s.page_width, s.page_height, s.top_margin, s.bottom_margin, s.left_margin, s.right_margin)


if __name__ == "__main__":
    main()
