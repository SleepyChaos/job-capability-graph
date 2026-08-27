import re
import zipfile
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document


DOCX = Path(r"D:\git\job-capability-graph\outputs\institution_database_report\具身智能机构库与企业产业链数据建设报告_比赛润色版_20260825.docx")
PDF = Path(r"D:\git\job-capability-graph\outputs\institution_database_report\qa_polished_v2\polished_v2.pdf")


def all_text(doc: Document) -> str:
    pieces = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            pieces.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        pieces.extend(p.text for p in section.header.paragraphs)
        pieces.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(pieces)


def main() -> None:
    doc = Document(DOCX)
    text = all_text(doc)
    banned = [
        "用户指定",
        "比赛材料推荐表述",
        "数据底座",
        "五层信息架构",
        "分类证据链",
        "机构画像",
        "能力画像",
        "循环搜索",
        "金标准",
        "在《具身智能机构库与企业产业链数据整理及动态更新验证方案",
    ]
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)} sections={len(doc.sections)}")
    print(f"pages={len(pdfium.PdfDocument(PDF))}")
    print("banned_hits=" + repr([term for term in banned if term in text]))
    expected = [
        "1,618条机构记录",
        "35个字段",
        "652条记录",
        "上游124条",
        "中游437条",
        "下游52条",
        "横向支撑39条",
        "不列示未经验证的准确率指标",
        "不作为当前已完成成果",
    ]
    print("missing_expected=" + repr([term for term in expected if term not in text]))
    with zipfile.ZipFile(DOCX) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        names = archive.namelist()
    print(f"tracked_insertions={len(re.findall(r'<w:ins(?:\s|>)', xml))} tracked_deletions={len(re.findall(r'<w:del(?:\s|>)', xml))}")
    print(f"comments_part={'word/comments.xml' in names}")
    print(f"title={doc.core_properties.title!r} subject={doc.core_properties.subject!r}")


if __name__ == "__main__":
    main()
