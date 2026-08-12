"""简历文件安全与多格式文本提取（总体设计 §10、后端设计 §12.1）。

P1 范围：
- MIME 嗅探（magic bytes），不完全信任客户端自报类型；
- 大小限制与扩展名白名单；
- 旧 DOC 隔离（不在主解析链路处理）；
- 病毒扫描接口预留（当前返回 skipped，等待真实引擎接入，见需确认清单 Q10）；
- 文本 PDF / DOCX / TXT 提取；扫描 PDF 明确报错降级（OCR 未接入）。
"""

from __future__ import annotations

import io
from pathlib import PurePosixPath

MAX_RESUME_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
QUARANTINED_EXTENSIONS = {".doc"}

INPUT_TYPE_BY_EXTENSION = {
    ".txt": "txt",
    ".pdf": "pdf_text",
    ".docx": "docx_text",
}


class ResumeFileError(ValueError):
    """A user-correctable resume file error."""


def sniff_mime(data: bytes) -> str:
    if data.startswith(b"%PDF"):
        return "application/pdf"
    if data.startswith(b"PK\x03\x04"):
        return (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return "text/plain"


def virus_scan(data: bytes) -> dict:
    """病毒扫描接口预留：真实引擎接入前返回 skipped（Q10）。"""
    return {"engine": "none", "status": "skipped", "size": len(data)}


def extract_resume_text(filename: str, data: bytes) -> tuple[str, str, str]:
    """返回 (文本, MIME, input_type_code)。"""
    if len(data) > MAX_RESUME_BYTES:
        raise ResumeFileError("简历文件超过10MB大小限制")
    if not data:
        raise ResumeFileError("简历文件为空")
    extension = PurePosixPath(filename).suffix.lower()
    if extension in QUARANTINED_EXTENSIONS:
        raise ResumeFileError("旧版 DOC 需走隔离转换流程，暂不支持直接解析，请粘贴文本")
    if extension not in ALLOWED_EXTENSIONS:
        raise ResumeFileError(f"不支持的文件类型：{extension or '未知'}（支持 txt/pdf/docx）")
    sniffed = sniff_mime(data)
    if extension == ".pdf":
        if sniffed != "application/pdf":
            raise ResumeFileError("文件内容与 PDF 扩展名不符，已拒绝解析")
        return _extract_pdf(data), sniffed, INPUT_TYPE_BY_EXTENSION[extension]
    if extension == ".docx":
        if not sniffed.startswith("application/vnd.openxmlformats"):
            raise ResumeFileError("文件内容与 DOCX 扩展名不符，已拒绝解析")
        return _extract_docx(data), sniffed, INPUT_TYPE_BY_EXTENSION[extension]
    return _extract_txt(data), sniffed, "txt"


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - 向用户返回可纠正错误
        raise ResumeFileError(f"PDF 解析失败：{exc}") from exc
    parts = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(part.strip() for part in parts if part.strip())
    if len(text.strip()) < 30:
        raise ResumeFileError(
            "未从 PDF 提取到文本（可能是扫描件）；OCR 尚未接入，请粘贴文本或上传文本型 PDF"
        )
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ResumeFileError(f"DOCX 解析失败：{exc}") from exc
    parts = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(part for part in parts if part)
    if len(text.strip()) < 30:
        raise ResumeFileError("DOCX 中未提取到有效文本")
    return text


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeFileError("文本编码无法识别（支持 UTF-8 / GBK）")
