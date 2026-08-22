"""把 4.1-4.2-6.1-7.1.md 渲染成竞赛用 .docx，并把 【插图】 占位替换为对应 PNG。

**为什么要有这个脚本。** 正文与插图会随实验结论反复修订，手工重排 Word 既慢又
容易漏掉某处数字。把渲染固化成脚本，正文改完重跑一次即可，图与文不会脱节。

**Markdown 子集。** 只支持正文实际用到的构造，不追求通用：
标题 #/##/###、段落、粗体 **、行内代码 `、行内公式 $…$、独立公式 $$…$$、
表格、无序/有序列表、引用块（其中 【插图】 会被替换成图片）。

**行内公式按记号逐个转换，不走 LaTeX 排版。** 竞赛稿里的公式都是单层的
（下标、上标、\\frac、\\sum、希腊字母），用 Word 自带的上下标格式足够表达，
引入 OMML 转换器会带来与本项目无关的依赖。转换不了的记号会显式报错而不是
静默输出原始 LaTeX——插图与公式里出现裸 \\frac 是最容易被漏看的一类错误。

用法（docs/competition 目录）：
    python3 build_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "4.1-4.2-6.1-7.1.md"
TARGET = HERE / "4.1-4.2-6.1-7.1.docx"
FIGURE_WIDTH = Cm(15.5)

BODY_FONT = "微软雅黑"
MONO_FONT = "Consolas"
MATH_FONT = "Cambria Math"

# LaTeX 记号 → 普通字符。只收录正文用到的，未收录的会触发报错。
SYMBOLS = {
    r"\theta": "θ", r"\lambda": "λ", r"\alpha": "α", r"\beta": "β",
    r"\times": "×", r"\cdot": "·", r"\in": "∈", r"\ge": "≥", r"\le": "≤",
    r"\geq": "≥", r"\leq": "≤", r"\neq": "≠", r"\cup": "∪", r"\cap": "∩",
    r"\subseteq": "⊆", r"\emptyset": "∅", r"\sum": "Σ", r"\max": "max",
    r"\min": "min", r"\exp": "exp", r"\log": "log", r"\left": "", r"\right": "",
    r"\{": "{", r"\}": "}", r"\,": " ", r"\;": " ", r"\quad": "　",
    r"\to": "→", r"\approx": "≈", r"\infty": "∞",
    # \text / \mathrm 只是把内容排成正体，去掉命令保留内容即可。
    r"\text": "", r"\mathrm": "", r"\Rightarrow": "⇒",
}


class RenderError(Exception):
    """渲染失败。宁可中断也不要产出带裸 LaTeX 或缺图的稿子。"""


# --------------------------------------------------------------------------- 公式

def math_tokens(expr: str) -> list[tuple[str, str]]:
    """把行内公式拆成 (文本, 样式) 序列，样式 ∈ {normal, sub, sup}。

    先展开 \\frac，再逐字符扫描 _{}/^{}，最后替换符号——顺序不能换：
    \\frac 的分子里常带下标，先处理下标会把 \\frac 的花括号配对打乱。
    """
    expr = expand_frac(expr)
    out: list[tuple[str, str]] = []
    buffer = ""
    index = 0
    while index < len(expr):
        char = expr[index]
        if char in "_^" and index + 1 < len(expr):
            style = "sub" if char == "_" else "sup"
            index += 1
            if expr[index] == "{":
                depth, start = 1, index + 1
                index += 1
                while index < len(expr) and depth:
                    if expr[index] == "{":
                        depth += 1
                    elif expr[index] == "}":
                        depth -= 1
                    index += 1
                inner = expr[start : index - 1]
            else:
                inner = expr[index]
                index += 1
            if buffer:
                out.append((substitute(buffer), "normal"))
                buffer = ""
            out.append((substitute(inner), style))
            continue
        buffer += char
        index += 1
    if buffer:
        out.append((substitute(buffer), "normal"))
    return [(text, style) for text, style in out if text]


def expand_frac(expr: str) -> str:
    r"""\frac{a}{b} → (a) / (b)。分子分母可能自带花括号，所以要配对扫描。"""
    while True:
        pos = expr.find(r"\frac")
        if pos < 0:
            return expr
        cursor = pos + len(r"\frac")
        parts = []
        for _ in range(2):
            while cursor < len(expr) and expr[cursor] == " ":
                cursor += 1
            if cursor >= len(expr) or expr[cursor] != "{":
                raise RenderError(rf"\frac 缺少参数：{expr}")
            depth, start = 1, cursor + 1
            cursor += 1
            while cursor < len(expr) and depth:
                if expr[cursor] == "{":
                    depth += 1
                elif expr[cursor] == "}":
                    depth -= 1
                cursor += 1
            parts.append(expr[start : cursor - 1])
        expr = f"{expr[:pos]}({parts[0]}) / ({parts[1]}){expr[cursor:]}"


def substitute(text: str) -> str:
    for token, replacement in sorted(SYMBOLS.items(), key=lambda kv: -len(kv[0])):
        text = text.replace(token, replacement)
    text = text.replace("{", "").replace("}", "")
    if "\\" in text:
        raise RenderError(f"公式里有未支持的 LaTeX 记号：{text!r}")
    return text


def add_math(paragraph, expr: str) -> None:
    for text, style in math_tokens(expr):
        run = paragraph.add_run(text)
        run.font.name = MATH_FONT
        run.font.italic = True
        if style == "sub":
            run.font.subscript = True
        elif style == "sup":
            run.font.superscript = True


# --------------------------------------------------------------------------- 行内

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)")


def add_inline(paragraph, text: str, bold: bool = False) -> None:
    """单趟切分行内标记。粗体内部可能再嵌公式或代码，因此递归处理。"""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            add_inline(paragraph, piece[2:-2], bold=True)
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(10)
            run.font.bold = bold
        elif piece.startswith("$") and piece.endswith("$") and len(piece) > 2:
            add_math(paragraph, piece[1:-1])
        else:
            run = paragraph.add_run(piece)
            run.font.bold = bold


# --------------------------------------------------------------------------- 块级

def flush_table(document, rows: list[str]) -> None:
    cells = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    header, body = cells[0], cells[2:]  # cells[1] 是对齐行
    table = document.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, text in enumerate(header):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        add_inline(paragraph, text, bold=True)
    for row_index, row in enumerate(body, start=1):
        for col_index, text in enumerate(row):
            if col_index >= len(header):
                continue
            add_inline(table.rows[row_index].cells[col_index].paragraphs[0], text)
    document.add_paragraph()


def add_figure(document, caption: str) -> None:
    name = caption.replace("【插图】", "").strip()
    number = name.split("　")[0].replace("图 ", "图")
    matches = sorted(HERE.glob(f"{number}_*.png"))
    if not matches:
        raise RenderError(f"找不到插图：{number}（正文引用了 {name}）")
    document.add_picture(str(matches[0]), width=FIGURE_WIDTH)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(name)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
        elif stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            flush_table(document, block)
        elif stripped == "$$":
            index += 1
            expr = []
            while index < len(lines) and lines[index].strip() != "$$":
                expr.append(lines[index].strip())
                index += 1
            index += 1
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_math(paragraph, " ".join(expr))
        elif stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            document.add_heading(stripped.lstrip("#").strip(), level=min(level, 4))
            index += 1
        elif stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            if "【插图】" in content:
                add_figure(document, content.strip("*"))
            else:
                paragraph = document.add_paragraph(style="Intense Quote")
                add_inline(paragraph, content)
            index += 1
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
        elif re.match(r"^\d+\.\s", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s", "", stripped))
            index += 1
        elif stripped == "---":
            index += 1
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, stripped)
            index += 1

    document.save(TARGET)
    print(f"已生成 {TARGET.name}")


if __name__ == "__main__":
    try:
        build()
    except RenderError as error:
        sys.exit(f"渲染失败：{error}")
