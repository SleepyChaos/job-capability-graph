from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "work"
OUT = WORK / "具身智能岗位与能力三图谱_赛题对齐与整体架构说明_评审版.docx"

SCREEN_ECOSYSTEM = WORK / "graph-v05-industry-role.png"
SCREEN_DISCOVERY = WORK / "graph-v05-technology-role.png"
SCREEN_PORTRAIT = WORK / "graph-v05-portrait.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
CYAN = "24A6A8"
ORANGE = "F28E2B"
GREEN = "2E8B57"
RED = "C65D4B"
INK = "243442"
MUTED = "5D6B78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
PALE_CYAN = "EAF7F6"
PALE_ORANGE = "FFF4E8"
PALE_GREEN = "EDF7F1"
PALE_RED = "FBEFEC"
LIGHT_GRAY = "F4F5F7"
MID_GRAY = "D6DDE4"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        data = kwargs[edge]
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        element = borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in data:
                element.set(qn(f"w:{key}"), str(data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    table_indent = tbl_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    if widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        dxa_widths = [round(width * 1440) for width in widths]
        for width in dxa_widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
                    tc_width = row.cells[idx]._tc.get_or_add_tcPr().get_or_add_tcW()
                    tc_width.set(qn("w:w"), str(dxa_widths[idx]))
                    tc_width.set(qn("w:type"), "dxa")


def set_run_font(run, size=11, bold=False, color=INK, name="Calibri") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_run(p, text: str, size=11, bold=False, color=INK, name="Calibri"):
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, name=name)
    return r


def fmt_para(p, before=0, after=6, line=1.10, keep=False, align=None) -> None:
    f = p.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing = line
    f.keep_with_next = keep
    f.widow_control = True
    if align is not None:
        p.alignment = align


def add_body(doc, text: str, *, bold_prefix: str | None = None, after=6, color=INK):
    p = doc.add_paragraph()
    fmt_para(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=NAVY)
        add_run(p, text[len(bold_prefix):], color=color)
    else:
        add_run(p, text, color=color)
    return p


def add_bullet(doc, text: str, level=0, color=INK, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["List Bullet"]
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    fmt_para(p, after=8, line=1.167)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=NAVY)
        add_run(p, text[len(bold_prefix):], color=color)
    else:
        add_run(p, text, color=color)
    return p


def add_heading(doc, text: str, level=1, *, number: str | None = None):
    p = doc.add_paragraph()
    if level == 1:
        fmt_para(p, before=16, after=8, keep=True)
        if number:
            add_run(p, number + "  ", size=12, bold=True, color=CYAN)
        add_run(p, text, size=16, bold=True, color=BLUE)
    elif level == 2:
        fmt_para(p, before=12, after=6, keep=True)
        add_run(p, text, size=13, bold=True, color=BLUE)
    else:
        fmt_para(p, before=8, after=4, keep=True)
        add_run(p, text, size=12, bold=True, color=DARK_BLUE)
    return p


def add_kicker(doc, text: str):
    p = doc.add_paragraph()
    fmt_para(p, after=8, keep=True)
    add_run(p, text.upper(), size=9, bold=True, color=CYAN)
    return p


def add_statement(doc, headline: str, detail: str = "", fill=PALE_BLUE, accent=BLUE):
    t = doc.add_table(rows=1, cols=1)
    set_table_fixed(t, [6.5])
    c = t.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_margins(c, top=140, bottom=140, start=180, end=180)
    set_cell_border(c, left={"val": "single", "sz": 24, "color": accent})
    p = c.paragraphs[0]
    fmt_para(p, after=3)
    add_run(p, headline, size=12, bold=True, color=NAVY)
    if detail:
        p2 = c.add_paragraph()
        fmt_para(p2, after=0)
        add_run(p2, detail, size=10.5, color=MUTED)
    return t


def add_metric_cards(doc, cards: list[tuple[str, str, str]], fills=None):
    t = doc.add_table(rows=1, cols=len(cards))
    widths = [6.5 / len(cards)] * len(cards)
    set_table_fixed(t, widths)
    fills = fills or [PALE_BLUE] * len(cards)
    for i, (num, label, note) in enumerate(cards):
        c = t.cell(0, i)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c, fills[i % len(fills)])
        set_cell_margins(c, top=120, bottom=110, start=90, end=90)
        set_cell_border(c, right={"val": "single", "sz": 6, "color": WHITE})
        p = c.paragraphs[0]
        fmt_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, num, size=20, bold=True, color=NAVY)
        p2 = c.add_paragraph()
        fmt_para(p2, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p2, label, size=9.5, bold=True, color=INK)
        if note:
            p3 = c.add_paragraph()
            fmt_para(p3, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_run(p3, note, size=8.5, color=MUTED)
    return t


def add_flow(doc, steps: list[tuple[str, str]], colors=None):
    n = len(steps)
    cols = n * 2 - 1
    t = doc.add_table(rows=1, cols=cols)
    widths = []
    box = (6.5 - 0.20 * (n - 1)) / n
    for i in range(cols):
        widths.append(box if i % 2 == 0 else 0.20)
    set_table_fixed(t, widths)
    colors = colors or [PALE_BLUE, PALE_CYAN, PALE_ORANGE, PALE_GREEN]
    for i, (title, sub) in enumerate(steps):
        c = t.cell(0, i * 2)
        set_cell_shading(c, colors[i % len(colors)])
        set_cell_margins(c, top=100, bottom=100, start=70, end=70)
        p = c.paragraphs[0]
        fmt_para(p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, title, size=9.5, bold=True, color=NAVY)
        if sub:
            p2 = c.add_paragraph()
            fmt_para(p2, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_run(p2, sub, size=7.8, color=MUTED)
        if i < n - 1:
            a = t.cell(0, i * 2 + 1)
            a.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p3 = a.paragraphs[0]
            fmt_para(p3, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_run(p3, "→", size=13, bold=True, color=CYAN)
    return t


def add_comparison_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=LIGHT_BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(t, widths)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        set_cell_shading(c, header_fill)
        set_cell_margins(c)
        set_cell_border(c, bottom={"val": "single", "sz": 8, "color": BLUE})
        p = c.paragraphs[0]
        fmt_para(p, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_run(p, h, size=9.5, bold=True, color=NAVY)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, value in enumerate(row):
            c = cells[j]
            set_cell_shading(c, WHITE if i % 2 == 0 else "FAFBFC")
            set_cell_margins(c)
            set_cell_border(c, bottom={"val": "single", "sz": 4, "color": MID_GRAY})
            p = c.paragraphs[0]
            fmt_para(p, after=0)
            add_run(p, value, size=9.2, color=INK)
    return t


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    fmt_para(p, before=3, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, text, size=8.5, color=MUTED)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = sec.header
    hp = header.paragraphs[0]
    fmt_para(hp, after=0)
    add_run(hp, "具身智能岗位与能力图谱 · 赛题对齐评审说明", size=8.5, bold=True, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    add_run(fp, "评审版  |  证据候选版 v0.5  |  ", size=8.5, color=MUTED)
    add_page_number(fp)


def page_cover(doc: Document) -> None:
    add_kicker(doc, "COMPETITION ALIGNMENT  ·  DATA → THREE GRAPHS → MATCH")
    p = doc.add_paragraph()
    fmt_para(p, before=18, after=12, keep=True)
    add_run(p, "具身智能岗位与能力三图谱\n整体架构与赛题对齐说明", size=28, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    fmt_para(p2, after=24)
    add_run(p2, "作品设计实现方案 · 评审版", size=19, bold=True, color=BLUE)

    add_statement(
        doc,
        "一句话结论",
        "以 4,655 条 v4 岗位事实为底座，以标准岗位作为三图公共桥梁：产业图谱回答“谁在招”，技术图谱回答“需要什么技术”，岗位画像图谱回答“这一类岗位做什么、怎样演化”；所有结论都能回到具体 JD。",
        fill=PALE_CYAN,
        accent=CYAN,
    )
    doc.add_paragraph()
    add_metric_cards(
        doc,
        [
            ("4,655", "岗位事实", "v4 统一主表"),
            ("96.1%", "企业映射覆盖", "4,472 / 4,655"),
            ("64.5%", "技术映射覆盖", "3,005 / 4,655"),
        ],
        fills=[PALE_BLUE, PALE_CYAN, PALE_GREEN],
    )
    doc.add_paragraph()
    add_heading(doc, "评委先看这三件事", level=2)
    add_bullet(doc, "数据不是“拼表”：先统一字段、ID、标签和证据，再并入 v4。", bold_prefix="数据不是“拼表”")
    add_bullet(doc, "三张主图不是三套孤岛；产业、技术与岗位画像都通过标准岗位连接到真实 JD。", bold_prefix="三张主图")
    add_bullet(doc, "岗位簇发现是岗位画像图谱内部的方法层，不再单独冒充一张业务主图。", bold_prefix="岗位簇发现")
    add_bullet(doc, "前端可以完整演示；后端未正式测出的 HDBSCAN、90%准确率和动态趋势不写成既成成果。", bold_prefix="发布边界")
    p3 = doc.add_paragraph()
    fmt_para(p3, before=22, after=0)
    add_run(p3, "数据口径：岗位信息 v4 及企业增强分析工作簿  |  版本日期：2026-08-26", size=9, color=MUTED)


def page_overview(doc: Document) -> None:
    add_heading(doc, "10 秒看懂：一套事实底座，三张业务图，一个验证闭环", number="01")
    add_statement(doc, "标准岗位是三张图的公共桥梁，具体 JD 是共同证据底座。", "评委从任何一张图进入，最终都能看到岗位、企业、技术要求和原始 JD，而不是只看到抽象节点。")
    doc.add_paragraph()
    add_flow(
        doc,
        [
            ("多源数据", "岗位 / 企业 / 技术 / 人才"),
            ("治理主数据", "ID / 别名 / 版本 / 证据"),
            ("三张主图", "产业 / 技术 / 岗位画像"),
            ("业务应用", "发现 / 演化 / 匹配"),
            ("反馈再治理", "专家审核 / 新快照"),
        ],
    )
    add_heading(doc, "三张图分别回答什么？", level=2)
    add_comparison_table(
        doc,
        ["图谱", "10秒问题", "从哪里起图", "最终落点"],
        [
            ["产业—岗位图谱", "谁在招？产业需求在哪里？", "企业库：产业链、细分领域、融资、地区、总部城市", "企业→标准岗位→具体岗位/JD"],
            ["技术—岗位图谱", "某项技术被哪些岗位需要？", "技术主数据：7个L1、43个L2、229个L3、1,872个L4", "L4技术词→标准岗位→具体岗位/JD"],
            ["岗位画像与演化图谱", "这一类岗位做什么、要求什么、怎样变化？", "6方向→17类别→类内岗位簇→标准岗位", "五维画像→版本差异→JD证据"],
        ],
        [1.25, 1.55, 2.25, 1.45],
    )
    add_heading(doc, "与赛题核心功能逐项对齐", level=2)
    add_comparison_table(
        doc,
        ["赛题要求", "本项目承接位置", "当前状态"],
        [
            ["新岗位发现与定义", "17类别内部发现候选岗位簇；输出标准岗位、职责、技能与场景", "候选流程可演示，HDBSCAN待正式复算"],
            ["既有岗位能力动态更新", "标准岗位五维画像版本对比：新增/删除/修改/必备与加分迁移", "数据结构已设计，需下一期同口径快照"],
            ["岗位全景与技能点下钻", "技术主数据L1–L4→岗位；岗位画像按6/17/簇/标准岗位下钻", "前端已实现"],
            ["人岗匹配、差距与路径", "简历证据→标准岗位画像→JD/企业精排→差距→学习任务→再匹配", "闭环可演示，90%需金标准评测"],
        ],
        [1.55, 3.45, 1.5],
    )


def page_sources(doc: Document) -> None:
    add_heading(doc, "数据从哪里来：事实源、主数据、参考标准分层管理", number="02")
    add_statement(doc, "不同来源承担不同角色：岗位表给事实，企业库给企业属性，技术词表给技术层级。", "文档中的设计描述只用于梳理口径；最终数字以可复算的数据文件和当前v4产出为准。", fill=PALE_ORANGE, accent=ORANGE)
    doc.add_paragraph()
    add_comparison_table(
        doc,
        ["层级", "文件 / 数据", "规模", "在项目中的角色"],
        [
            ["核心事实源", "3700+条的清洗过的岗位数据.xlsx", "3,718 条", "第一版稳定基线；26 字段、唯一 ID、6/17 标签、清洗 JD、来源字段"],
            ["补充事实源", "后续采集岗位", "851 条原始 / 720 条入v4", "补 ID、统一类别、去重和质量审核后并入"],
            ["补充事实源", "图谱定向候选岗位", "220 条原始 / 217 条入v4", "标准化字段和标签后并入；3条未形成独立事实"],
            ["企业主数据", "具身智能企业数据_整合去重_完整.xlsx", "633 家", "提供产业链、细分领域、融资、地区、总部城市"],
            ["标准岗位词包", "搜索词包_按岗位.csv", "107 岗 / 610 变体", "定义标准岗位与真实标题别名；人工校准所属岗位簇"],
            ["技术主数据", "技术词主数据_20260727.xlsx", "7 / 43 / 229 / 1,872", "提供L1–L4技术层级，是技术图谱唯一分类口径"],
            ["参考标准", "具身智能从业人员能力要求（征求意见稿）", "参考文本", "用于术语与能力维度对照，不直接改写用户要求"],
        ],
        [1.0, 2.25, 0.85, 2.4],
    )
    add_heading(doc, "来源可信度不是“一刀切”", level=2)
    add_bullet(doc, "岗位表保留岗位名称、公司、JD、URL、来源与有效性字段，图谱可回到原始证据。")
    add_bullet(doc, "企业字段只从企业主数据表补充；多候选不自动选，未匹配保持为空并进入待补全表。")
    add_bullet(doc, "标准岗位名称来自搜索词包；标题变体只用于候选归并，还要同时通过 17 类 / 岗位簇边界和置信阈值。")
    add_bullet(doc, "AI/规则产生的候选标签单独标识，不能冒充人工确认的正式分类。")
    add_body(doc, "当前v4构成：核心基线3,718条 + 后续采集清洗入库720条 + 候选清洗入库217条 = 4,655条独立岗位事实。851和220是原始来源规模，不等于最终新增数。", color=MUTED)


def page_baseline(doc: Document) -> None:
    add_heading(doc, "为什么 3,718 条可作基线，而 851 + 220 不能直接拼接？", number="03")
    add_statement(doc, "差别不在“多少”，在“是否已经形成一致的事实口径”。", "3,718 条具备稳定结构和可追溯证据；另外两批需要先经过同一套治理门槛。")
    doc.add_paragraph()
    add_comparison_table(
        doc,
        ["检查项", "3,718 条核心数据", "851 条补充数据", "220 条补充数据"],
        [
            ["字段结构", "固定 26 字段", "原始字段不齐、格式混合", "24 字段，与 v4 不完全一致"],
            ["唯一标识", "已有唯一 occ_id", "没有统一 occ_id", "需重新分配并校验 ID"],
            ["职业标签", "已有 6 方向 / 17 类别", "出现非标准或混合类别", "需映射到最终 17 类"],
            ["JD 与证据", "已有清洗 JD、URL、来源、有效性", "重复、空值和噪声需处理", "文本较整齐，但仍需同口径校验"],
            ["可直接进入图谱", "可以：作为 v1 基线", "不可以：先清洗治理", "不可以：先对齐映射"],
        ],
        [1.15, 1.75, 1.75, 1.85],
    )
    add_heading(doc, "合并后发生了什么？", level=2)
    add_flow(
        doc,
        [
            ("3,718", "核心基线"),
            ("+ 851 + 220", "补充来源"),
            ("4,789", "采集记录"),
            ("清洗 / 去重", "映射 / 审核"),
            ("4,655", "v4 独立事实"),
        ],
        colors=[PALE_BLUE, PALE_ORANGE, PALE_ORANGE, PALE_CYAN, PALE_GREEN],
    )
    doc.add_paragraph()
    add_metric_cards(
        doc,
        [
            ("O = 3,718", "基线记录", "全部保留"),
            ("V = 937", "新增记录", "720 + 217 标准化后入库"),
            ("134", "未独立入库", "重复 / 冲突 / 归并"),
        ],
        fills=[PALE_BLUE, PALE_GREEN, PALE_RED],
    )
    add_body(doc, "审计复算显示，851 条来源中按“公司 + 岗位名称 + JD”归一化后有 744 个唯一签名，即至少 107 条重复签名；其余差异在字段标准化、有效性与标签冲突审核中被排除或吸收。这里说的是“没有形成新的独立事实”，不是简单删除。", color=MUTED)


def page_governance(doc: Document) -> None:
    add_heading(doc, "数据怎么处理、怎么交叉验证：每一步都有可核对的数", number="04")
    add_statement(doc, "处理链只做三件事：统一、去重、留证。", "每个自动结果都设置审计出口；无法确认的企业和字段留空，不为了完整率制造事实。", fill=PALE_CYAN, accent=CYAN)
    doc.add_paragraph()
    add_flow(
        doc,
        [
            ("字段对齐", "26 字段口径"),
            ("ID 治理", "O / V 唯一编码"),
            ("文本清洗", "标题 / JD / 公司"),
            ("标签映射", "6 方向 / 17 类"),
            ("去重与有效性", "证据签名"),
            ("企业增强", "可审计匹配"),
        ],
    )
    add_heading(doc, "八层交叉验证", level=2)
    add_comparison_table(
        doc,
        ["验证层", "可核对公式 / 规则", "当前结果"],
        [
            ["行数守恒", "3,718 + 851 + 220 = 4,789；独立事实 = 4,655", "差异 134 条有治理解释"],
            ["主键唯一", "O=3,718；V=937；occ_id 不重复", "4,655 个唯一岗位事实"],
            ["层级守恒", "各方向 / 类别 / 岗位簇岗位数汇总均回到 v4", "6 方向、17 类别、42 候选簇"],
            ["关系约束", "每个岗位只归属一个类别边界内岗位簇", "不允许跨类别自由漂移"],
            ["标准岗闸门", "人工层级 + 去通用职级词后的标题相似度 + 候选差值", "258 条高置信；4,397 条待专家映射"],
            ["企业审计", "已匹配 + 待补全 = 全量岗位", "4,472 + 183 = 4,655"],
            ["技术审计", "只用JD全文一致回接或L4精确命中；不做公司名+岗位名模糊猜测", "3,005条已关联；1,650条待补"],
            ["证据回溯", "标准画像点至少有 2 条 JD 支撑；可打开具体 JD / URL", "单条 JD 不直接形成标准结论"],
        ],
        [1.2, 3.5, 1.8],
    )
    add_heading(doc, "企业匹配的安全顺序", level=2)
    add_flow(
        doc,
        [
            ("人工别名", "优先"),
            ("全称精确", "一致即匹配"),
            ("别名精确", "别名表"),
            ("核心名", "唯一才通过"),
            ("唯一包含", "仍需唯一"),
            ("待补全", "多候选 / 未匹配"),
        ],
        colors=[PALE_GREEN, PALE_GREEN, PALE_CYAN, PALE_BLUE, PALE_ORANGE, PALE_RED],
    )
    add_body(doc, "产出工作簿同时保留“岗位信息v4_企业增强、企业匹配审计、企业待补全、图谱节点、图谱关系、企业库快照、字段与口径”等工作表，评委可以从结果一路追到口径。", color=MUTED)


def page_ecosystem(doc: Document) -> None:
    add_heading(doc, "图谱一｜产业—岗位图谱：企业库起图，岗位作为需求落点", number="05")
    add_statement(doc, "这张图回答“谁在招、产业需求在哪里”。", "企业属性以企业库为唯一口径；岗位数量和要求来自v4岗位事实。两者通过治理后的企业实体关联，不用JD反推融资、地区或产业链属性。")
    doc.add_paragraph()
    doc.add_picture(str(SCREEN_ECOSYSTEM), width=Inches(6.42))
    add_caption(doc, "系统实图：产业链层级→企业→标准岗位；右侧逐条查看具体岗位与JD")
    add_flow(
        doc,
        [
            ("企业库", "633条主数据"),
            ("企业属性", "产业链/领域/融资/地区/城市"),
            ("企业实体", "名称与别名治理"),
            ("标准岗位", "企业需求聚合"),
            ("具体岗位 / JD", "右侧证据"),
        ],
        colors=[PALE_BLUE, PALE_CYAN, PALE_GREEN, PALE_GREEN, PALE_ORANGE],
    )
    add_heading(doc, "当前事实与展示边界", level=2)
    add_bullet(doc, "4,472条岗位已关联102个企业实体，企业属性覆盖96.1%；183条保持待补全或待核验。")
    add_bullet(doc, "页面可按产业链层级、公司细分领域、融资轮次、所属地区、总部城市五个维度切换。")
    add_bullet(doc, "企业节点继续下钻到标准岗位；主图无需铺开每条JD，右侧证据区保留所有具体岗位。")
    add_bullet(doc, "产业图谱不负责定义岗位能力；它负责提供企业和产业背景，并把需求送到岗位画像图谱。")


def page_discovery(doc: Document) -> None:
    add_heading(doc, "图谱二｜技术—岗位图谱：技术主数据起图，尾端必须有岗位", number="06")
    add_statement(doc, "这张图回答“某项技术被哪些岗位需要”。", "技术分类只认L1–L4主数据；普通技能标签不能冒充技术体系。技术词通过可追溯证据连接标准岗位与具体JD。", fill=PALE_CYAN, accent=CYAN)
    doc.add_paragraph()
    doc.add_picture(str(SCREEN_DISCOVERY), width=Inches(6.42))
    add_caption(doc, "系统实图：L1→L2→L3→L4技术词→标准岗位→具体岗位/JD")
    add_flow(
        doc,
        [
            ("7个L1域", "技术方向"),
            ("43个L2类", "技术类别"),
            ("229个L3点", "标准技术点"),
            ("1,872个L4词", "表面词/术语"),
            ("标准岗位/JD", "业务落点"),
        ],
        colors=[PALE_BLUE, PALE_BLUE, PALE_CYAN, PALE_GREEN, PALE_ORANGE],
    )
    add_heading(doc, "技术词怎样安全落到v4岗位", level=2)
    add_comparison_table(
        doc,
        ["关联方式", "规则", "当前结果"],
        [
            ["JD全文一致回接", "旧技术标注的清洗JD与v4清洗JD全文归一化后一致", "2,403条岗位"],
            ["L4技术词精确命中", "v4技能标签与L4主数据标准词完全一致", "1,113条岗位"],
            ["合并去重后", "两种安全关系取并集；同一岗位不重复计算", "3,005条，覆盖64.5%"],
            ["待补区", "不使用公司名+岗位名模糊猜测，也不强行映射", "1,650条岗位"],
        ],
        [1.35, 3.4, 1.75],
    )
    add_body(doc, "技术图谱的创新点不是“画出很多术语”，而是把技术主数据、标准岗位与真实招聘证据连成可反向查询的链：既能从岗位看技术，也能从技术看岗位。", color=MUTED)


def page_portrait(doc: Document) -> None:
    add_heading(doc, "图谱三｜岗位画像与演化图谱：共性在图上，个体JD在右侧", number="07")
    add_statement(doc, "这张图回答“这一类岗位做什么、需要什么、怎样演化”。", "图谱v1提供技术与术语坐标，搜索词包定义标准岗位与名称变体，v4同类JD提供市场证据；三者共同形成画像。", fill=PALE_GREEN, accent=GREEN)
    doc.add_paragraph()
    doc.add_picture(str(SCREEN_PORTRAIT), width=Inches(6.42))
    add_caption(doc, "系统实图：标准岗位—五维标准点—具体岗位/完整JD证据")
    add_flow(
        doc,
        [
            ("6方向/17类别", "人工边界"),
            ("42候选岗位簇", "类别内发现"),
            ("107标准岗位", "610名称变体"),
            ("多JD五维画像", "至少2条支撑"),
            ("版本差异/JD", "演化与回溯"),
        ],
        colors=[PALE_CYAN, PALE_BLUE, PALE_GREEN, PALE_ORANGE, PALE_BLUE],
    )
    add_heading(doc, "岗位簇在这张图中的正确位置", level=2)
    add_bullet(doc, "岗位簇只在17个职业类别内部发现；当前42个为透明规则候选，目标流程是Embedding+HDBSCAN+专家校准。")
    add_bullet(doc, "标准岗位不是“代表岗位”，而是可承接一组名称变体和多条JD的业务实体；例如同类CEO/总裁标题归并后再生成五维画像。")
    add_bullet(doc, "五维包括职责、专业技能、通用能力、应用场景、任职条件；每个画像点显示支持JD数和覆盖率。")
    add_bullet(doc, "动态演化必须比较同口径时间快照，输出新增、删除、修改及必备↔加分迁移。当前只有v4单一快照，因此只发布演化基线，不伪造趋势。")


def page_matching(doc: Document) -> None:
    add_heading(doc, "人岗匹配｜不是一个分数，而是一条可验证的闭环", number="08")
    add_statement(doc, "系统先问“证据够不够”，再问“匹不匹配”。", "缺少某项能力只标记为证据不足；只有用户确认或有评测证据，才判定为真实缺口。")
    doc.add_paragraph()
    add_flow(
        doc,
        [
            ("简历 / 补充问答", "人才证据"),
            ("能力标准化", "L3 技术能力"),
            ("岗位簇召回", "方向 + 语义"),
            ("标准岗位匹配", "五维画像 + 目标"),
            ("JD / 企业精排", "具体机会 + 偏好"),
            ("差距解释", "双方证据"),
            ("学习路径", "任务 + 验证"),
            ("再匹配", "新画像版本"),
        ],
        colors=[PALE_BLUE, PALE_CYAN, PALE_GREEN, PALE_GREEN, PALE_ORANGE, PALE_BLUE, PALE_CYAN],
    )
    add_heading(doc, "当前 P0 可回放公式", level=2)
    add_metric_cards(
        doc,
        [
            ("40%", "必需能力覆盖", "岗位要求为分母"),
            ("10%", "可迁移能力", "同能力域相邻证据"),
            ("45%", "求职目标语义", "岗位名称与目标重合"),
            ("5%", "证据完整度", "标准能力数量有界得分"),
        ],
        fills=[PALE_GREEN, PALE_CYAN, PALE_BLUE, PALE_ORANGE],
    )
    add_body(doc, "总分由机械规则生成，LLM 不直接给分。第一层用方向 / 类别 / 岗位簇缩小范围，第二层按标准岗位五维画像匹配，第三层用具体 JD 与企业偏好精排。每个结果保存分项得分、权重、算法版本、具体 JD、差距证据与路径绑定；相同数据快照可以复算。")
    add_heading(doc, "企业库如何参与匹配", level=2)
    add_comparison_table(
        doc,
        ["层次", "当前用法", "边界"],
        [
            ["岗位适配", "能力与求职目标产生核心匹配分", "不使用性别、年龄、婚育等敏感属性"],
            ["企业偏好", "按产业链、细分领域、融资、地区、总部城市筛选和解释", "v0.4 不强行并入核心分；缺失字段不扣分"],
            ["闭环更新", "完成实践任务并提交证据，生成新画像版本后重新匹配", "“点完成”本身不会自动加分"],
        ],
        [1.2, 3.25, 2.05],
    )
    add_body(doc, "当前系统尚未建立专家标注的简历—岗位金标准集，因此可以说明“闭环已打通、分数可回放”，但不能宣称匹配准确率达到 90%。", color=RED)


def page_innovation(doc: Document) -> None:
    add_heading(doc, "创新点｜创新不在“用了 AI”，而在“AI 被放在正确的位置”", number="09")
    add_comparison_table(
        doc,
        ["创新", "传统做法容易出现的问题", "本项目的处理"],
        [
            ["1. 规则约束 + 语义发现 + 专家校准", "自由聚类跨职业混杂；纯人工分类缺少新发现", "6/17 先定边界，AI 只发现细粒度簇，专家命名与发布"],
            ["2. 标准岗位桥接“共性与证据”", "直接用一条JD画像会受单一企业影响；全量JD上图又不可读", "主图止于标准岗位；五维由多JD聚合，右侧保留具体JD"],
            ["3. 企业增强不污染岗位事实", "为了补全而猜公司属性", "确定性匹配；多候选/未匹配进入审计和待补全"],
            ["4. 匹配是可解释闭环", "只给一个推荐分，无法复核也无法提升", "分项得分→证据→差距→学习任务→新版本→再匹配"],
        ],
        [1.6, 2.2, 2.7],
    )
    add_heading(doc, "评委高频问题：一句话回答", level=2)
    add_comparison_table(
        doc,
        ["可能追问", "建议回答"],
        [
            ["为什么不是 17 个岗位簇？", "17 是职业类别边界；岗位簇是类别内部的细粒度发现层，解决同类岗位仍然差异很大的问题。"],
            ["为什么不直接全量聚类？", "全量自由聚类会被通用词干扰并跨职业混杂；先分类再聚类，结果更稳定、也更能解释。"],
            ["HDBSCAN 已经跑完了吗？", "没有。当前 42 个是透明规则候选层；目标流程是 Embedding + HDBSCAN + 专家校准，系统已明确标注。"],
            ["为什么图上不放全部 4,655 条 JD？", "主图要让人看懂，所以止于标准岗位；原始岗位没有删除，用户可在右侧逐条查看完整 JD。"],
            ["五维画像是不是一条 JD 生成的？", "不是。画像以标准岗位为中心，每个标准点至少由 2 条同类 JD 支撑，并显示数量和覆盖率。"],
            ["为什么只有 258 条已映射？", "当前采用高置信发布闸门，宁可把 4,397 条留给专家，也不因“工程师/经理”等通用词错配。"],
            ["为什么是 4,655 而不是 4,789？", "4,789 是采集行，4,655 是治理后的独立事实；134 行因重复、冲突或归并没有新增事实。"],
            ["企业信息缺失怎么办？", "留空并进入待补全，不用 AI 猜；当前 4,472 条已匹配，183 条待审。"],
            ["匹配分可靠吗？", "当前可复算、可解释，但尚未完成金标准准确率验证；所以不夸大准确率，只展示闭环和证据。"],
        ],
        [1.65, 4.85],
    )


def page_release(doc: Document) -> None:
    add_heading(doc, "结论与发布边界｜评委可以相信什么，下一步验证什么", number="10")
    add_statement(doc, "当前版本已经能稳定展示“数据—图谱—匹配”闭环。", "最重要的可信点是：结果有口径、数字能对账、节点可追证、算法有边界。", fill=PALE_GREEN, accent=GREEN)
    doc.add_paragraph()
    add_comparison_table(
        doc,
        ["状态", "内容", "对外表述"],
        [
            ["已完成", "4,655 条岗位 v4；6/17/42 层级；107 标准岗位、610 标题变体；三图谱；企业增强审计；P0 匹配闭环", "可以现场演示和复核"],
            ["候选可用", "62 个标准岗位已有JD证据；258 条通过高置信映射；企业字段覆盖 96.1%", "证据充分者展示画像，缺失明确留空"],
            ["下一阶段", "4,397 条JD专家映射；类别内 Embedding + HDBSCAN；簇稳定性、人工命名与金标准评测", "作为路线，不写成既成成果"],
        ],
        [1.15, 3.55, 1.8],
    )
    add_heading(doc, "下一轮验证计划", level=2)
    add_flow(
        doc,
        [
            ("冻结 v4", "数据 / 词表 / 代码版本"),
            ("类别内向量化", "标题 + JD + 技能"),
            ("HDBSCAN", "参数 / 噪声 / 稳定性"),
            ("候选对照", "规则簇 vs 语义簇"),
            ("专家校准", "名称 / 合并 / 拆分"),
            ("金标准评测", "准确率分项报告"),
        ],
        colors=[PALE_BLUE, PALE_CYAN, PALE_CYAN, PALE_ORANGE, PALE_GREEN, PALE_GREEN],
    )
    add_heading(doc, "一分钟答辩话术", level=2)
    add_statement(
        doc,
        "“我们不是让 AI 替代职业分类，而是让它在人工标准内发现更细的岗位结构。”",
        "“项目先把 4,789 条采集记录治理为 4,655 条可追溯岗位事实，再用 6 大方向和 17 类别约束岗位簇发现，形成 42 个可解释候选簇。搜索词包的 107 个标准岗位承接 610 个标题变体：主图止于标准岗位，五维画像由同类多条 JD 共同形成，具体招聘岗位与完整 JD 在右侧逐条核验。当前 258 条 JD 通过高置信闸门，其余保留专家映射，不为追求覆盖率制造错配。匹配端给出证据、差距和学习路径；企业信息只作可审计增强，缺失不猜。”",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_heading(doc, "最终方法表述", level=2)
    add_body(doc, "本研究采用“规则约束 + 语义聚类 + 专家校准”的岗位发现方法。首先利用 6 个职业方向和 17 个职业类别构建边界，再基于岗位名称、JD、公司与技能信息在类别内部发现细粒度岗位簇；随后以搜索词包定义标准岗位和名称变体，经过人工层级校准与标题置信闸门归并多条 JD。五维画像以标准岗位为中心，画像点至少由 2 条 JD 共同支持，具体 JD 保留为可回溯证据。人岗匹配依次经过岗位簇召回、标准岗位画像匹配、具体 JD / 企业精排、差距解释与学习反馈，从而形成可解释、可回放、可迭代的闭环。")
    p = doc.add_paragraph()
    fmt_para(p, before=12, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p, "—— 评审版方法说明结束 ——", size=9, color=MUTED)


def build() -> Path:
    for p in (SCREEN_ECOSYSTEM, SCREEN_DISCOVERY, SCREEN_PORTRAIT):
        if not p.exists():
            raise FileNotFoundError(p)
    doc = Document()
    configure_document(doc)
    page_cover(doc)
    for page_fn in (
        page_overview,
        page_sources,
        page_baseline,
        page_governance,
        page_ecosystem,
        page_discovery,
        page_portrait,
        page_matching,
        page_innovation,
        page_release,
    ):
        add_page_break(doc)
        page_fn(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
